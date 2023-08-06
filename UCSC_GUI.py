import PySimpleGUI as sg
import os
from window_layout import layout
from docx import Document

window = sg.Window('UCSC lncRNA Profile', layout=layout, icon=r'assets/bio.ico', size=(400, 500))

lnc_dict = {}


def lnc_class():
    if values['SENSE']:
        lnc_type = 'sense'
    elif values['ANTI-SENSE']:
        lnc_type = 'anti-sense'
    elif values['BIDIRECTIONAL']:
        lnc_type = 'bidirectional'
    elif values['INTRON']:
        lnc_type = 'intron'
    elif values['INTERGENIC']:
        lnc_type = 'intergenic'
    else:
        lnc_type = 'NULL'

    return lnc_type


def get_Epigenetics():
    if values['EA']:
        epi_type = 'Epigenetically Activated'
    elif values['ES']:
        epi_type = 'Epigenetically Silenced'
    else:
        epi_type = 'Unknown'
    return epi_type


def get_expression():
    expressions = []
    for tissue in ['BREAST', 'PANCREAS', 'SKIN', 'PROSTATE', 'LIVER', 'TESTIS', 'BLADDER', 'LUNG', 'INTESTINE', 'COLON',
                   'UTERUS', 'BRAIN', 'KIDNEY', 'CERVIX', 'STOMACH', 'THYROID']:
        if values[tissue]:
            expressions.append(tissue)
    e_sc = values['Expression_SC']
    return expressions, e_sc


def get_regulation():
    pro_position = values['P_POSITION']
    pro_id = values['P_ID']
    pro_sc = values['PROMOTER_SC']
    prolifs = {"Proliferation Data": [pro_position, pro_id, pro_sc]}

    TSS_position = values['TSS_POSITION']
    TSS_strand = values['TSS_STRAND']
    TSS_id = values['TSS_ID']
    TSS_sc = values['TSS_SC']
    TSSs = {"TSSs": [TSS_position, TSS_strand, TSS_id, TSS_sc]}

    CpG_position = values['POSITION']
    CpG_size = values['SIZE']
    CpG_count = values['CPG_CNT']
    CpG_sc = values['CpG_SC']
    CpGs = {"CpG Islands": [CpG_position, CpG_size, CpG_count, CpG_sc]}

    TAF = values['TAFs']
    TAF_sc = values['TAF_SC']
    TAFs = {"TAFs": [TAF, TAF_sc]}

    return prolifs, TSSs, CpGs, TAFs


def write_to_file():
    folder_name = "data/" + values['lnc_NAME']
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    else:
        sg.popup_yes_no("Are You Sure You Want To Override " + values['lnc_NAME'], title="Override Error")

    file_path = os.path.join(folder_name, f"{values['lnc_NAME']}_output.txt")
    with open(file_path, 'w') as file:
        for key, value in lnc_dict.items():
            file.write(f"{key}: {value}\n")

    sg.popup_auto_close("File Saved!")


def compile_data():
    doc = Document()
    sub_dicts = []

    for key, value in lnc_dict.items():
        if isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    sub_dicts.append(item)

    Gene_Data = sub_dicts[0]
    Gene_Data = Gene_Data['Gene Data']

    Epigenetics = sub_dicts[1]

    Tissue_Expression = sub_dicts[2]
    Tissue_Expression = Tissue_Expression['Tissues with Expression']

    Proliferation = sub_dicts[3]
    Proliferation = Proliferation['Proliferation Data']
    TSS = sub_dicts[4]
    TSS = TSS['TSSs']
    CpG_Islands = sub_dicts[5]
    CpG_Islands = CpG_Islands['CpG Islands']
    TAF = sub_dicts[6]
    TAF = TAF['TAFs']

    first_paragraph = f"Gene Name: {values['lnc_NAME']}, ENSEMBL Gene ID: {Gene_Data[0][1]}, ENSEMBL Transcript ID: " \
                      f"{Gene_Data[0][0]}, \n Gene Start: {Gene_Data[1][0]}, Gene End: {Gene_Data[1][1]}, Chromosome: " \
                      f"{Gene_Data[2][0]}, Band: {Gene_Data[2][1]}," \
                      f"Strand: {Gene_Data[3]}"
    first_paragraph_image = Gene_Data[4]

    second_paragraph = ""
    for key in Epigenetics.keys():
        second_paragraph = f"{values['lnc_NAME']} is {key}, with an expression value of {Epigenetics[key]}"

    third_paragraph = f"{values['lnc_NAME']} is expressed in {Tissue_Expression[0]}"
    third_paragraph_image = Tissue_Expression[1]

    fourth_paragraph = f'{values["lnc_NAME"]} has a promoter at position {Proliferation[0]} and is found with the ' \
                       f' Promoter ID: {Proliferation[1]}'
    fourth_paragraph_image = Proliferation[2]

    fifth_paragraph = f'{values["lnc_NAME"]} has a TSS at position {TSS[0]} on strand {TSS[1]} and is found with the' \
                      f' TSS ID: {TSS[2]}'
    fifth_paragraph_image = TSS[3]

    sixth_paragraph = f'{values["lnc_NAME"]} a promoter with TAF motifs for: {TAF[0]}'
    sixth_paragraph_image = TAF[1]

    seventh_paragraph = f'{values["lnc_NAME"]} has the nearest CpG relative to the promoter at ' \
                        f'a position of {CpG_Islands[0]} with a size of {CpG_Islands[1]}bp and a CpG Count ' \
                        f'of {CpG_Islands[2]}'
    seventh_paragraph_image = CpG_Islands[3]

    doc.add_paragraph(first_paragraph)

    try:
        doc.add_picture(first_paragraph_image)
    except FileNotFoundError:
        pass

    doc.add_paragraph(second_paragraph)
    doc.add_paragraph(third_paragraph)

    try:
        doc.add_picture(third_paragraph_image)
    except FileNotFoundError:
        pass

    doc.add_paragraph(fourth_paragraph)
    try:
        doc.add_picture(fourth_paragraph_image)
    except FileNotFoundError:
        pass

    doc.add_paragraph(fifth_paragraph)
    try:
        doc.add_picture(fifth_paragraph_image)
    except FileNotFoundError:
        pass

    doc.add_paragraph(sixth_paragraph)
    try:
        doc.add_picture(sixth_paragraph_image)
    except FileNotFoundError:
        pass

    doc.add_paragraph(seventh_paragraph)
    try:
        doc.add_picture(seventh_paragraph_image)
    except FileNotFoundError:
        pass

    output = values['OUTPUT']
    doc.save(f"{output}/{values['lnc_NAME']}.docx")


while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED or event == 'Quit':
        break

    if event == 'Submit':
        lnc_type = lnc_class()
        metadata = {"Gene Data": [
            [values['lnc_ENSEMBL_T'], values['lnc_ENSEMBL']], [values['lnc_START'], values['lnc_END']],
            [values['lnc_CHR'], values['lnc_BAND']], lnc_type, values['GTEx_SC']
        ]}

        epi_type = get_Epigenetics()
        epidata = {epi_type: values['EXP_VAL']}

        expressions, e_sc = get_expression()
        expressiondata = {"Tissues with Expression": [expressions, e_sc]}

        prolifs, TSSs, CpGs, TAFs = get_regulation()
        lnc_dict[values['lnc_NAME']] = [metadata, epidata, expressiondata, prolifs, TSSs, CpGs, TAFs]

        window.close()
        write_to_file()
        compile_data()

window.close()
